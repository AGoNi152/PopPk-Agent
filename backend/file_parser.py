from __future__ import annotations

import io
from pathlib import Path
from typing import Iterable

import fitz
import pandas as pd
from docx import Document
from fastapi import UploadFile

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt", ".md"}

async def extract_text_from_upload(file: UploadFile) -> str:
    name = file.filename or "uploaded_file"
    suffix = Path(name).suffix.lower()
    content = await file.read()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {suffix or 'unknown'}")
    if suffix == ".pdf":
        return extract_pdf(content, name)
    if suffix == ".docx":
        return extract_docx(content, name)
    if suffix in {".xlsx", ".xls"}:
        return extract_excel(content, name)
    if suffix == ".csv":
        return extract_csv(content, name)
    return extract_plain_text(content, name)

def extract_pdf(content: bytes, name: str) -> str:
    doc = fitz.open(stream=content, filetype="pdf")
    pages = []
    for index, page in enumerate(doc, start=1):
        pages.append(f"\n--- {name} | page {index} ---\n{page.get_text('text')}")
    return "\n".join(pages).strip()

def extract_docx(content: bytes, name: str) -> str:
    document = Document(io.BytesIO(content))
    parts = [f"--- {name} | DOCX text ---"]
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"\n--- {name} | table {table_index} ---")
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts).strip()

def extract_excel(content: bytes, name: str) -> str:
    sheets = pd.read_excel(io.BytesIO(content), sheet_name=None, dtype=str)
    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"\n--- {name} | sheet: {sheet_name} ---")
        preview = df.fillna("").head(200)
        parts.append(preview.to_csv(index=False))
    return "\n".join(parts).strip()

def extract_csv(content: bytes, name: str) -> str:
    df = pd.read_csv(io.BytesIO(content), dtype=str)
    return f"--- {name} | CSV preview ---\n{df.fillna('').head(300).to_csv(index=False)}"

def extract_plain_text(content: bytes, name: str) -> str:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("latin-1", errors="ignore")
    return f"--- {name} | text ---\n{text}"

def combine_materials(text: str | None, extracted: Iterable[str]) -> str:
    sections = []
    if text and text.strip():
        sections.append(f"--- User pasted text ---\n{text.strip()}")
    for item in extracted:
        if item and item.strip():
            sections.append(item.strip())
    return "\n\n".join(sections).strip()
