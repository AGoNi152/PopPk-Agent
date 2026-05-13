from __future__ import annotations

import uuid
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

OUTPUT_DIR = Path(__file__).parent / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

def create_docx(review: dict[str, Any]) -> str:
    filename = f"review_{uuid.uuid4().hex}.docx"
    path = OUTPUT_DIR / filename
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    document.add_heading("PopPK/ER Evidence Review Report", level=1)
    document.add_heading("Decision-readiness rating", level=2)
    document.add_paragraph(review.get("readiness", ""))
    document.add_heading("Executive summary", level=2)
    document.add_paragraph(review.get("executiveSummary", ""))

    document.add_heading("Issue tracker", level=2)
    for issue in review.get("issues", []):
        document.add_heading(f"{issue.get('severity', 'Issue')}: {issue.get('title', '')}", level=3)
        document.add_paragraph(issue.get("detail", ""))
        document.add_paragraph(f"Recommended action: {issue.get('action', '')}")

    document.add_heading("Completeness checklist", level=2)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Status"
    for row in review.get("checklist", []):
        cells = table.add_row().cells
        cells[0].text = row.get("item", "")
        cells[1].text = row.get("status", "")

    document.add_heading("Potential regulatory questions", level=2)
    for question in review.get("regulatoryQuestions", []):
        document.add_paragraph(question, style="List Number")

    document.add_heading("Recommended additional analyses", level=2)
    for item in review.get("additionalAnalyses", []):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Suggested regulatory-style language", level=2)
    document.add_paragraph(review.get("suggestedLanguage", ""))

    document.add_heading("Beginner explanation", level=2)
    document.add_paragraph(review.get("beginnerExplanation", ""))

    document.save(path)
    return filename

def create_pdf(review: dict[str, Any]) -> str:
    filename = f"review_{uuid.uuid4().hex}.pdf"
    path = OUTPUT_DIR / filename
    styles = getSampleStyleSheet()
    story = []

    def add_heading(text: str) -> None:
        story.append(Paragraph(text, styles["Heading2"]))
        story.append(Spacer(1, 8))

    def add_para(text: str) -> None:
        safe = str(text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe, styles["BodyText"]))
        story.append(Spacer(1, 8))

    story.append(Paragraph("PopPK/ER Evidence Review Report", styles["Title"]))
    story.append(Spacer(1, 12))
    add_heading("Decision-readiness rating")
    add_para(review.get("readiness", ""))
    add_heading("Executive summary")
    add_para(review.get("executiveSummary", ""))
    add_heading("Issue tracker")
    for issue in review.get("issues", []):
        add_para(f"{issue.get('severity', 'Issue')}: {issue.get('title', '')}")
        add_para(issue.get("detail", ""))
        add_para(f"Recommended action: {issue.get('action', '')}")
    add_heading("Completeness checklist")
    for row in review.get("checklist", []):
        add_para(f"{row.get('item', '')}: {row.get('status', '')}")
    add_heading("Potential regulatory questions")
    for question in review.get("regulatoryQuestions", []):
        add_para(f"• {question}")
    add_heading("Recommended additional analyses")
    for item in review.get("additionalAnalyses", []):
        add_para(f"• {item}")
    add_heading("Suggested regulatory-style language")
    add_para(review.get("suggestedLanguage", ""))
    add_heading("Beginner explanation")
    add_para(review.get("beginnerExplanation", ""))

    SimpleDocTemplate(str(path), pagesize=A4).build(story)
    return filename
